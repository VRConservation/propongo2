"""Export functionality for PDF, HTML, and DOCX generation."""

import os
import logging
import re as _re
from io import BytesIO
from typing import Tuple
from weasyprint import HTML
from flask import Blueprint, render_template, request, send_file, jsonify, Response
from markupsafe import Markup
from .models import Proposal
from .utils import build_export_context, build_tracker_export_context
from .config import Config, ERROR_MESSAGES

logger = logging.getLogger(__name__)

export_bp = Blueprint("export", __name__)

EXPORT_DIR = Config.EXPORTS_DIR


def ensure_export_dir() -> None:
    """Ensure export directory exists."""
    os.makedirs(EXPORT_DIR, exist_ok=True)


def _md_to_plain(text: str) -> str:
    """Strip markdown to plain text for DOCX."""
    if not text:
        return ""
    text = _re.sub(r'#{1,6}\s+', '', text)
    text = _re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = _re.sub(r'\*(.+?)\*', r'\1', text)
    text = _re.sub(r'`(.+?)`', r'\1', text)
    text = _re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    text = _re.sub(r'^[-*]\s+', '  - ', text, flags=_re.MULTILINE)
    return text.strip()


def _build_proposal_docx(proposal) -> BytesIO:
    """Build a DOCX document for a proposal."""
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(proposal.title or "Proposal")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)

    if proposal.client_name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Funder: {proposal.client_name}")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    if proposal.subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Program: {proposal.subtitle}")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(proposal.updated_at[:10] if proposal.updated_at else "")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()

    if proposal.project_summary:
        doc.add_heading('Project Summary', level=1)
        for line in _md_to_plain(proposal.project_summary).split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip())

    if proposal.scope:
        doc.add_heading('Scope', level=1)
        for line in _md_to_plain(proposal.scope).split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip())

    if proposal.budget_items:
        doc.add_heading('Budget', level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ['Item', 'Cost/Unit', 'Units', 'Total']
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)

        for task in proposal.tasks:
            task_items = [b for b in proposal.budget_items if b.get("task_id") == task["id"]]
            if not task_items:
                continue
            task_total = sum(i.get("cost_per_unit", 0) * i.get("units", 0) for i in task_items)
            row = table.add_row()
            row.cells[0].text = task["name"]
            row.cells[3].text = f"{task_total:,.0f}"
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(10)
            for item in task_items:
                total = item.get("cost_per_unit", 0) * item.get("units", 0)
                row = table.add_row()
                row.cells[0].text = f"  {item['name']}"
                row.cells[1].text = f"{item.get('cost_per_unit', 0):,.0f}"
                row.cells[2].text = f"{int(item.get('units', 0))}"
                row.cells[3].text = f"{total:,.0f}"
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)

        row = table.add_row()
        row.cells[0].text = 'Subtotal'
        row.cells[3].text = f"{proposal.total_budget:,.0f}"
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        indirect_percent = getattr(proposal, 'indirect_percent', 0) or 0
        if indirect_percent > 0:
            indirect_amount = proposal.total_budget * (indirect_percent / 100)
            row = table.add_row()
            row.cells[0].text = f"Indirect ({indirect_percent:.0f}%)"
            row.cells[3].text = f"{indirect_amount:,.0f}"

        row = table.add_row()
        row.cells[0].text = 'Total'
        indirect_amount = proposal.total_budget * (indirect_percent / 100)
        row.cells[3].text = f"${proposal.total_budget + indirect_amount:,.0f}"
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

    if proposal.show_budget_description and proposal.budget_description:
        doc.add_heading('Budget Description', level=1)
        for line in _md_to_plain(proposal.budget_description).split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip())

    if proposal.qualifications:
        doc.add_heading('Qualifications', level=1)
        for line in _md_to_plain(proposal.qualifications).split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip())

    if proposal.custom_sections:
        for section in sorted(proposal.custom_sections, key=lambda s: s.get("order", 0)):
            doc.add_heading(section.get("title", "Section"), level=1)
            for line in _md_to_plain(section.get("content", "")).split('\n'):
                if line.strip():
                    doc.add_paragraph(line.strip())

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _build_tracker_docx(proposal, ctx) -> BytesIO:
    """Build a DOCX document for the project tracker."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Project Report: {proposal.title}")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)

    if proposal.client_name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{proposal.client_name}{(' - ' + proposal.subtitle) if proposal.subtitle else ''}")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    doc.add_heading('Dashboard Summary', level=1)
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['Tasks Completed', 'Total Tasks', 'Total Budget', 'Milestones']
    values = [f"{ctx['overall_pct']}%", str(len(ctx['tasks'])), f"${ctx['total_with_indirect']:,.0f}", str(len(ctx['milestones']))]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
    for i, v in enumerate(values):
        table.rows[1].cells[i].text = v
        for run in table.rows[1].cells[i].paragraphs[0].runs:
            run.font.size = Pt(10)

    status_labels = {
        "not_started": "Not Started",
        "in_progress": "In Progress",
        "completed": "Completed",
        "delayed": "Delayed",
    }

    doc.add_heading('Tasks', level=1)
    for task in ctx['tasks']:
        doc.add_heading(task.get('name', 'Untitled'), level=2)
        status = status_labels.get(task.get('status', 'not_started'), 'Not Started')
        progress = task.get('progress_pct', 0)
        p = doc.add_paragraph()
        run = p.add_run(f"Status: {status}")
        run.bold = True
        p = doc.add_paragraph(f"Progress: {progress}%")
        if task.get('actual_start'):
            p = doc.add_paragraph(f"Actual Start: {task['actual_start']}")
        if task.get('actual_end'):
            p = doc.add_paragraph(f"Actual End: {task['actual_end']}")
        if task.get('notes'):
            doc.add_paragraph()
            run = p.add_run("Notes:")
            run.bold = True
            for line in task['notes'].split('\n'):
                if line.strip():
                    doc.add_paragraph(line.strip())

    doc.add_heading('Budget Tracking', level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['Task', 'Item', 'Planned', 'Actual', 'Variance']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)

    for task in ctx['tasks']:
        tid = task.get("id", "")
        if tid not in ctx['task_budgets']:
            continue
        tb = ctx['task_budgets'][tid]
        row = table.add_row()
        row.cells[0].text = task.get('name', '')
        row.cells[2].text = f"${tb['subtotal']:,.0f}"
        row.cells[3].text = f"${tb['actual_total']:,.0f}"
        variance = tb['actual_total'] - tb['subtotal']
        row.cells[4].text = f"${variance:,.0f}"
        for cell in row.cells:
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(9)
        for item in tb["items"]:
            planned = item.get("cost_per_unit", 0) * item.get("units", 0)
            actual = item.get("actual_cost", 0)
            v = actual - planned
            row = table.add_row()
            row.cells[1].text = item.get('name', '')
            row.cells[2].text = f"${planned:,.0f}"
            row.cells[3].text = f"${actual:,.0f}"
            row.cells[4].text = f"${v:,.0f}"
            for cell in row.cells:
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(9)

    row = table.add_row()
    row.cells[0].text = 'Total'
    row.cells[2].text = f"${ctx['total_with_indirect']:,.0f}"
    row.cells[3].text = f"${ctx['total_actual'] + ctx['indirect_amount']:,.0f}"
    for cell in row.cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)

    if ctx['milestones']:
        doc.add_heading('Milestones', level=1)
        for m in ctx['milestones']:
            check = "[x]" if m.get("completed") else "[ ]"
            date_str = f" ({m['date']})" if m.get('date') else ''
            doc.add_paragraph(f"{check} {m.get('name', '')}{date_str}")

    if ctx['reports']:
        doc.add_heading('Reports', level=1)
        for r in reversed(ctx['reports']):
            doc.add_heading(r.get('title', 'Report'), level=2)
            if r.get('date'):
                p = doc.add_paragraph(f"Date: {r['date']}")
                p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            if r.get('content'):
                for line in r['content'].split('\n'):
                    if line.strip():
                        doc.add_paragraph(line.strip())

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@export_bp.route("/export/pdf/<proposal_id>")
def export_pdf(proposal_id: str) -> Tuple[Response, int] | Response:
    """Export proposal as PDF."""
    proposal = Proposal.load(proposal_id)
    if not proposal:
        logger.warning(f"Proposal not found for PDF export: {proposal_id}")
        return jsonify(ERROR_MESSAGES['PROPOSAL_NOT_FOUND']), 404

    ctx = build_export_context(proposal)
    html_content = render_template("export_proposal.html", **ctx)

    ensure_export_dir()
    pdf_path = os.path.join(EXPORT_DIR, f"{proposal_id}.pdf")
    HTML(string=html_content, base_url=request.host_url).write_pdf(pdf_path)

    return send_file(
        pdf_path,
        mimetype="application/pdf",
        as_attachment=True,
        download_name=f"{proposal.title or 'proposal'}.pdf",
    )


@export_bp.route("/export/html/<proposal_id>")
def export_html(proposal_id: str) -> Tuple[Response, int] | Response:
    """Export proposal as HTML file."""
    proposal = Proposal.load(proposal_id)
    if not proposal:
        logger.warning(f"Proposal not found for HTML export: {proposal_id}")
        return jsonify(ERROR_MESSAGES['PROPOSAL_NOT_FOUND']), 404

    ctx = build_export_context(proposal)
    html_content = render_template("export_proposal.html", **ctx)

    return Response(
        html_content,
        mimetype="text/html",
        headers={"Content-Disposition": "inline"},
    )


@export_bp.route("/export/docx/<proposal_id>")
def export_docx(proposal_id: str) -> Tuple[Response, int] | Response:
    """Export proposal as DOCX file."""
    proposal = Proposal.load(proposal_id)
    if not proposal:
        logger.warning(f"Proposal not found for DOCX export: {proposal_id}")
        return jsonify(ERROR_MESSAGES['PROPOSAL_NOT_FOUND']), 404

    try:
        buf = _build_proposal_docx(proposal)
    except ImportError:
        return jsonify({"error": "python-docx not installed"}), 500
    except Exception as e:
        logger.error(f"DOCX export failed: {e}")
        return jsonify({"error": f"DOCX export failed: {str(e)}"}), 500

    return send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=f"{proposal.title or 'proposal'}.docx",
    )


@export_bp.route("/export/tracker/pdf/<proposal_id>")
def export_tracker_pdf(proposal_id: str) -> Tuple[Response, int] | Response:
    """Export tracker as PDF."""
    proposal = Proposal.load(proposal_id)
    if not proposal:
        return jsonify(ERROR_MESSAGES['PROPOSAL_NOT_FOUND']), 404

    ctx = build_tracker_export_context(proposal)
    html_content = render_template("export_tracker.html", **ctx)

    ensure_export_dir()
    pdf_path = os.path.join(EXPORT_DIR, f"tracker_{proposal_id}.pdf")
    HTML(string=html_content, base_url=request.host_url).write_pdf(pdf_path)

    return send_file(
        pdf_path,
        mimetype="application/pdf",
        as_attachment=True,
        download_name=f"{proposal.title or 'project'}_tracker.pdf",
    )


@export_bp.route("/export/tracker/html/<proposal_id>")
def export_tracker_html(proposal_id: str) -> Tuple[Response, int] | Response:
    """Export tracker as HTML."""
    proposal = Proposal.load(proposal_id)
    if not proposal:
        return jsonify(ERROR_MESSAGES['PROPOSAL_NOT_FOUND']), 404

    ctx = build_tracker_export_context(proposal)
    html_content = render_template("export_tracker.html", **ctx)

    return Response(
        html_content,
        mimetype="text/html",
        headers={"Content-Disposition": "inline"},
    )


@export_bp.route("/export/tracker/docx/<proposal_id>")
def export_tracker_docx(proposal_id: str) -> Tuple[Response, int] | Response:
    """Export tracker as DOCX."""
    proposal = Proposal.load(proposal_id)
    if not proposal:
        return jsonify(ERROR_MESSAGES['PROPOSAL_NOT_FOUND']), 404

    ctx = build_tracker_export_context(proposal)

    try:
        buf = _build_tracker_docx(proposal, ctx)
    except ImportError:
        return jsonify({"error": "python-docx not installed"}), 500
    except Exception as e:
        logger.error(f"Tracker DOCX export failed: {e}")
        return jsonify({"error": f"DOCX export failed: {str(e)}"}), 500

    return send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=f"{proposal.title or 'project'}_tracker.docx",
    )
